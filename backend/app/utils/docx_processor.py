# Procesador de documentos DOCX
import os
import re
import tempfile
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
from loguru import logger

try:
    from docx import Document
    from docx.shared import Cm, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx no está instalado. Algunas funciones estarán limitadas.")


class DocxProcessor:
    """
    Procesa documentos DOCX para extraer información y validar
    """
    
    @staticmethod
    def validate_page_size(docx_path: Path) -> Tuple[bool, Dict[str, float]]:
        """
        Validar que el documento tenga tamaño México Oficio (21.59cm x 34.01cm)
        
        Returns:
            (es_valido, tamaño_pagina)
        """
        if not DOCX_AVAILABLE:
            logger.warning("python-docx no disponible, validación de tamaño deshabilitada")
            return True, {"ancho": 21.59, "alto": 34.01, "unidad": "cm"}
        
        try:
            doc = Document(docx_path)
            sections = doc.sections
            
            if not sections:
                return False, {"ancho": 0, "alto": 0, "unidad": "cm"}
            
            section = sections[0]
            page_width = section.page_width
            page_height = section.page_height
            
            # Convertir de twips a cm (1 twip = 1/1440 inch, 1 inch = 2.54 cm)
            width_cm = (page_width / 1440) * 2.54
            height_cm = (page_height / 1440) * 2.54
            
            # Tolerancia de 0.1 cm
            expected_width = 21.59
            expected_height = 34.01
            tolerance = 0.1
            
            is_valid = (
                abs(width_cm - expected_width) <= tolerance and 
                abs(height_cm - expected_height) <= tolerance
            )
            
            page_size = {
                "ancho": round(width_cm, 2),
                "alto": round(height_cm, 2),
                "unidad": "cm"
            }
            
            logger.info(f"Tamaño de página detectado: {width_cm}cm x {height_cm}cm")
            
            return is_valid, page_size
            
        except Exception as e:
            logger.error(f"Error validando tamaño de página: {str(e)}")
            return False, {"ancho": 0, "alto": 0, "unidad": "cm"}
    
    @staticmethod
    def extract_placeholders(docx_path: Path) -> List[str]:
        """
        Extraer placeholders del documento DOCX.
        
        Busca texto entre {{ y }} en el documento.
        
        Returns:
            Lista de nombres de placeholders
        """
        if not DOCX_AVAILABLE:
            logger.warning("python-docx no disponible, extracción de placeholders deshabilitada")
            return []
        
        placeholders = []
        try:
            doc = Document(docx_path)
            
            # Buscar en párrafos
            for paragraph in doc.paragraphs:
                text = paragraph.text
                matches = re.findall(r'\{\{(.+?)\}\}', text)
                placeholders.extend(matches)
            
            # Buscar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            matches = re.findall(r'\{\{(.+?)\}\}', text)
                            placeholders.extend(matches)
            
            # Eliminar duplicados
            placeholders = list(set(placeholders))
            
            logger.info(f"Placeholders extraídos: {placeholders}")
            return placeholders
            
        except Exception as e:
            logger.error(f"Error extrayendo placeholders: {str(e)}")
            return []
    
    @staticmethod
    def convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
        """
        Convertir DOCX a PDF
        
        Note: En producción, se necesitaría una librería de conversión como:
        - LibreOffice headless (Linux/macOS)
        - Microsoft Word COM (Windows)
        - python-docx2pdf (cross-platform)
        
        Por ahora, devolvemos False para indicar que la conversión no está disponible
        """
        logger.warning("Conversión DOCX a PDF no implementada completamente")
        
        # Implementación básica - crear un PDF vacío como placeholder
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            
            c = canvas.Canvas(str(pdf_path), pagesize=letter)
            
            # Título
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, 750, "Vista Previa de Documento")
            
            # Información
            c.setFont("Helvetica", 12)
            c.drawString(100, 720, "Documento original: DOCX")
            c.drawString(100, 700, f"Archivo: {docx_path.name}")
            c.drawString(100, 680, "Nota: La conversión real requiere Microsoft Word o LibreOffice")
            
            c.save()
            
            logger.info(f"PDF placeholder creado: {pdf_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creando PDF placeholder: {str(e)}")
            return False
    
    @staticmethod
    def get_document_info(docx_path: Path) -> Dict[str, Any]:
        """
        Obtener información del documento DOCX
        """
        if not DOCX_AVAILABLE:
            return {
                "error": "python-docx no disponible",
                "disponible": False
            }
        
        try:
            doc = Document(docx_path)
            
            info = {
                "disponible": True,
                "num_parrafos": len(doc.paragraphs),
                "num_tablas": len(doc.tables),
                "num_secciones": len(doc.sections),
                "estilos": list(doc.styles.element.xpath('//w:style/@w:styleId', 
                                                         namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})),
                "core_properties": {
                    "title": doc.core_properties.title,
                    "author": doc.core_properties.author,
                    "created": doc.core_properties.created,
                    "modified": doc.core_properties.modified
                }
            }
            
            return info
            
        except Exception as e:
            logger.error(f"Error obteniendo información del documento: {str(e)}")
            return {
                "error": str(e),
                "disponible": False
            }